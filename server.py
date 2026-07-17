import os
import sys
import shutil
import math
import zipfile
import pickle
import glob
import json
import traceback
import subprocess
import io
from copy import copy
from datetime import datetime
from typing import Dict, Any, List
import openpyxl
from openpyxl.drawing.image import Image as OpenpyxlImage
from openpyxl.chart import ScatterChart, Reference, Series
from openpyxl.chart.axis import ChartLines
from openpyxl.chart.shapes import GraphicalProperties
from openpyxl.drawing.line import LineProperties
from openpyxl.drawing.spreadsheet_drawing import TwoCellAnchor
from openpyxl.chart.text import RichText
from openpyxl.drawing.text import Paragraph, ParagraphProperties, CharacterProperties, Font as DrawFont
from openpyxl.drawing.colors import ColorChoice
from openpyxl.chart.layout import Layout, ManualLayout
from fastapi import FastAPI, HTTPException, status, UploadFile, File, Request
from fastapi.responses import StreamingResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field

import re
import pandas as pd
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

def load_dotenv_manually():
    """
    Manually parses a local .env file if it exists and loads key-value pairs
    into os.environ. This avoids adding a dependency on python-dotenv.
    """
    import os
    env_path = ".env"
    if os.path.exists(env_path):
        try:
            with open(env_path, "r", encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if not line or line.startswith("#"):
                        continue
                    if "=" in line:
                        key, val = line.split("=", 1)
                        key = key.strip()
                        val = val.strip().strip("'\"")
                        os.environ[key] = val
            print(" [INFO] Loaded environment variables from .env manually.")
        except Exception as e:
            print(f" [WARNING] Failed to read .env file manually: {e}")

load_dotenv_manually()

# =====================================================================
# Configuration & Constants
# =====================================================================
ZIP_NAME = "soilaiapp-main.zip"
EXTRACT_DIR = "soilaiapp-main"
REQUIRED_FILES = [
    "RF_soil_classifier.pkl",
    "RF_gi_regressor.pkl",
    "soil_classifier_label_encoder.pkl",
    "cbr_regressor_demo.pkl"
]

app = FastAPI(
    title="Geotechnical Soil AI API Server",
    description="FastAPI Local Bridge Backend for soil classification, CBR, and Group Index prediction.",
    version="2.0.0"
)

# Enable CORS for frontend integration
app.add_middleware(
    CORSMiddleware,
    allow_origin_regex=r"https?://.*",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Global variables to store the loaded models
models: Dict[str, Any] = {}
model_status: Dict[str, str] = {}
using_ml_models = False

# Global cache for dynamic ML model swap & calibration
tf_model = None
joblib_scaler = None
calibrations_cache = {}

# =====================================================================
# 1. Extraction & Verification
# =====================================================================
def extract_and_verify_archive():
    """
    Checks if soilaiapp-main.zip exists in the root folder.
    Safely unzips its contents and verifies the integrity of the files.
    """
    print("\n" + "="*60)
    print(" [PHASE 1] ARCHIVE EXTRACTION & VERIFICATION ")
    print("="*60)
    
    zip_path = ZIP_NAME
    
    # Check if the zip file exists
    if not os.path.exists(zip_path):
        print(f" [WARNING] '{zip_path}' not found in the root directory.")
        print(f"           Checking if '{EXTRACT_DIR}' already exists...")
        if os.path.exists(EXTRACT_DIR):
            print(f" [INFO] Extracted directory '{EXTRACT_DIR}' exists. Verifying contents...")
            check_extracted_files()
            return
        else:
            print(" [ERROR] Neither the ZIP archive nor the extracted folder was found!")
            print("         The server will operate in high-fidelity USCS Fallback Mode.")
            return

    # Check if we need to extract
    needs_extraction = False
    if not os.path.exists(EXTRACT_DIR):
        needs_extraction = True
    else:
        # Check if all required files exist inside
        for f in REQUIRED_FILES:
            if not os.path.exists(os.path.join(EXTRACT_DIR, f)):
                needs_extraction = True
                print(f" [INFO] Missing required file '{f}'. Re-extracting archive...")
                break

    if needs_extraction:
        print(f" [INFO] Extracting '{zip_path}' to '{EXTRACT_DIR}'...")
        try:
            os.makedirs(EXTRACT_DIR, exist_ok=True)
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall(".")
            print(" [SUCCESS] Extraction completed successfully.")
        except Exception as e:
            print(f" [ERROR] Failed to unzip archive: {e}")
            print("         The server will operate in high-fidelity USCS Fallback Mode.")
            return
    else:
        print(" [INFO] Existing extraction folder verified. Skipping unzip.")

    check_extracted_files()


def check_extracted_files():
    """Verifies that key machine learning models are present."""
    missing = []
    print(" [INFO] Verifying extracted files:")
    for f in REQUIRED_FILES:
        full_path = os.path.join(EXTRACT_DIR, f)
        if os.path.exists(full_path):
            print(f"   - [OK]  {f} ({os.path.getsize(full_path)} bytes)")
        else:
            print(f"   - [MISSING] {f}")
            missing.append(f)
            
    if missing:
        print(f" [WARNING] Missing files: {missing}. Some ML functionalities might be limited.")
    else:
        print(" [SUCCESS] All core geotechnical model files verified successfully.")
    print("="*60 + "\n")


# =====================================================================
# 2. Model Initialization & Logging
# =====================================================================
def initialize_models():
    """
    Loads all pickled geotechnical models into memory.
    Provides detailed terminal logging of the initialization process.
    """
    global models, using_ml_models, model_status
    print("="*60)
    print(" [PHASE 2] GEOTECHNICAL ML PIPELINE INITIALIZATION ")
    print("="*60)
    
    # Try importing required ML libraries. If missing, we use USCS fallback.
    try:
        import sklearn
        import numpy as np
        print(f" [INFO] Scikit-learn detected (v{sklearn.__version__}). Attempting to load pickles...")
    except ImportError:
        print(" [WARNING] Scikit-learn or NumPy is not installed in the Python environment.")
        print("           Models cannot be loaded. Using high-fidelity Geotechnical USCS Rule Engine.")
        using_ml_models = False
        print("="*60 + "\n")
        return

    # Attempt to load each model file
    files_to_load = {
        "rf_classifier": "RF_soil_classifier.pkl",
        "xgb_classifier": "XGB_soil_classifier.pkl",
        "lgbm_classifier": "LGBM_soil_classifier.pkl",
        "rf_gi_regressor": "RF_gi_regressor.pkl",
        "xgb_gi_regressor": "XGB_gi_regressor.pkl",
        "label_encoder": "soil_classifier_label_encoder.pkl"
    }

    loaded_count = 0
    
    for key, filename in files_to_load.items():
        path = os.path.join(EXTRACT_DIR, filename)
        if not os.path.exists(path):
            model_status[key] = "File Missing"
            print(f" [LOAD FAILED] {filename:35s} | File not found.")
            continue
            
        try:
            with open(path, 'rb') as f:
                models[key] = pickle.load(f)
            model_status[key] = "Loaded Successfully"
            loaded_count += 1
            print(f" [LOAD OK]     {filename:35s} | Model loaded into memory.")
        except Exception as e:
            model_status[key] = f"Error: {str(e)}"
            print(f" [LOAD FAILED] {filename:35s} | Error: {e}")

    # Load CBR Regressor (search for the latest one)
    cbr_pattern = os.path.join(EXTRACT_DIR, "cbr_regressor_*.pkl")
    cbr_files = glob.glob(cbr_pattern)
    
    cbr_file_to_load = None
    if cbr_files:
        # Sort files to find the latest full trained model or fallback to demo
        cbr_files.sort()
        # Prefer full models over demo if available
        full_models = [f for f in cbr_files if "full" in f]
        if full_models:
            cbr_file_to_load = full_models[-1] # latest full
        else:
            cbr_file_to_load = os.path.join(EXTRACT_DIR, "cbr_regressor_demo.pkl")
    else:
        # Check if direct demo exists
        demo_path = os.path.join(EXTRACT_DIR, "cbr_regressor_demo.pkl")
        if os.path.exists(demo_path):
            cbr_file_to_load = demo_path

    if cbr_file_to_load:
        try:
            with open(cbr_file_to_load, 'rb') as f:
                models["cbr_regressor"] = pickle.load(f)
            model_status["cbr_regressor"] = f"Loaded ({os.path.basename(cbr_file_to_load)})"
            loaded_count += 1
            print(f" [LOAD OK]     {os.path.basename(cbr_file_to_load):35s} | CBR Model loaded into memory.")
        except Exception as e:
            model_status["cbr_regressor"] = f"Error: {str(e)}"
            print(f" [LOAD FAILED] {os.path.basename(cbr_file_to_load):35s} | Error: {e}")
    else:
        model_status["cbr_regressor"] = "File Missing"
        print(f" [LOAD FAILED] cbr_regressor_*.pkl                 | No matching file found.")

    if loaded_count > 0:
        using_ml_models = True
        print(f"\n [SUCCESS] Geotechnical ML Pipeline ready. {loaded_count} models online.")
    else:
        using_ml_models = False
        print("\n [WARNING] No models could be loaded. Defaulting to Geotechnical USCS Rule Engine.")
    print("="*60 + "\n")


# =====================================================================
# 3. High-Fidelity Geotechnical Fallback (USCS & Empirical Rules)
# =====================================================================
def run_uscs_fallback_engine(
    PL: float, PI: float, D10: float, D30: float, D60: float, Cu: float, Cc: float, OMC: float, MDD: float
) -> Dict[str, Any]:
    """
    Geotechnical USCS (Unified Soil Classification System) and empirical analysis engine.
    Calculates classification, CBR estimation, and Group Index based on soil mechanics formulas.
    """
    LL = PL + PI
    
    # 1. Determine Soil Classification using USCS rules
    soil_code = "CL"
    soil_name = "Lean Clay (Fine-grained soil)"
    description = ""
    
    # We estimate sieve #200 passing percentage (Fines %) based on grain size.
    # D10 is the size where 10% of particles are smaller.
    # Sieve #200 is 0.075 mm.
    # If D10 < 0.075, then at least 10% passes sieve #200.
    # If D30 < 0.075, then at least 30% passes.
    # If D60 < 0.075, then at least 60% passes.
    if D10 < 0.075:
        fines_percent = 65.0 # High fine content estimation
    else:
        fines_percent = 12.0 # Coarse-grained with some fines
        
    # Classification Logic
    if PI >= 7 and PI >= 0.73 * (LL - 20):
        # Above A-Line, Plastic
        if LL < 50:
            soil_code = "CL"
            soil_name = "Lean Clay"
            description = "Low to medium plasticity clay, typical of glacial tills, alluvial clays, or clayey soils."
        else:
            soil_code = "CH"
            soil_name = "Fat Clay"
            description = "High plasticity inorganic clay, highly expansive, low permeability, poor subgrade material."
    elif PI < 4 or PI < 0.73 * (LL - 20):
        # Below A-Line, Silt
        if LL < 50:
            soil_code = "ML"
            soil_name = "Silt"
            description = "Low plasticity inorganic silt, high frost susceptibility, low cohesion."
        else:
            soil_code = "MH"
            soil_name = "Elastic Silt"
            description = "High plasticity silt, micaceous or diatomaceous, high elasticity and compressibility."
    else:
        # Transition Zone (Dual classification)
        soil_code = "CL-ML"
        soil_name = "Silty Clay"
        description = "Transition soil containing significant mixtures of clay and silt with low-plasticity features."

    # If the grain sizes indicate sand or gravel (very coarse)
    if D10 > 0.15:
        if Cu >= 6 and (1.0 <= Cc <= 3.0):
            soil_code = "SW"
            soil_name = "Well-graded Sand"
            description = "Clean, well-graded sand with wide range of particle sizes. Excellent foundation/subgrade material."
        elif Cu < 6 or Cc < 1.0 or Cc > 3.0:
            soil_code = "SP"
            soil_name = "Poorly-graded Sand"
            description = "Uniform sand or gap-graded sand, clean, low cohesion, moderate drainage."
        elif Cu >= 4 and (1.0 <= Cc <= 3.0):
            soil_code = "GW"
            soil_name = "Well-graded Gravel"
            description = "Clean, well-graded gravel-sand mixtures. Outstanding load-bearing capacity."
        else:
            soil_code = "GP"
            soil_name = "Poorly-graded Gravel"
            description = "Clean, uniform gravel, highly permeable, good drainage but poor grading."

    # 2. AASHTO Group Index (GI) Calculation
    # GI = (Fines - 35) * [0.2 + 0.005 * (LL - 40)] + 0.01 * (Fines - 15) * (PI - 10)
    f_term = max(0.0, fines_percent - 35.0)
    ll_term = max(0.0, LL - 40.0)
    f_term_pi = max(0.0, fines_percent - 15.0)
    pi_term = max(0.0, PI - 10.0)
    
    gi_value = (f_term * (0.2 + 0.005 * ll_term)) + (0.01 * f_term_pi * pi_term)
    gi_value = max(0.0, gi_value) # GI is never negative
    
    # 3. Empirical California Bearing Ratio (CBR) Calculation (%)
    # CBR is heavily dependent on Soil Type, MDD, and OMC.
    if soil_code in ["CL", "CH"]:
        # Fine-grained clays have low CBR (2 - 8%)
        base_cbr = 3.0
        cbr_est = base_cbr + 2.5 * (MDD - 1.4) - 0.08 * PI - 0.1 * abs(OMC - 16.0)
        cbr_est = max(1.5, min(8.0, cbr_est))
    elif soil_code in ["ML", "MH"]:
        # Silts have slightly better CBR (4 - 12%)
        base_cbr = 5.0
        cbr_est = base_cbr + 3.0 * (MDD - 1.5) - 0.05 * PI - 0.08 * abs(OMC - 15.0)
        cbr_est = max(2.5, min(12.0, cbr_est))
    elif soil_code in ["SW", "SP", "CL-ML"]:
        # Sands have moderate to good CBR (12 - 35%)
        base_cbr = 18.0
        cbr_est = base_cbr + 15.0 * (MDD - 1.7) + 1.2 * Cu - 0.15 * OMC
        cbr_est = max(10.0, min(38.0, cbr_est))
    else: # GW / GP
        # Gravels have high CBR (35 - 80%)
        base_cbr = 45.0
        cbr_est = base_cbr + 25.0 * (MDD - 1.9) + 2.0 * Cu
        cbr_est = max(35.0, min(85.0, cbr_est))

    # Formulate recommendations based on CBR
    if cbr_est < 5.0:
        compaction_rec = "Very poor subgrade. Stabilization with lime or cement, or subgrade replacement/geogrid reinforcement is highly recommended."
    elif cbr_est < 10.0:
        compaction_rec = "Fair subgrade. Subgrade needs controlled moisture compaction. Good compaction at OMC (+/- 2%) is vital."
    elif cbr_est < 30.0:
        compaction_rec = "Good subgrade / base. Standard compaction to 95% MDD is sufficient. Suitable for sub-base courses."
    else:
        compaction_rec = "Excellent subgrade / base. Outstanding bearing capacity. Ideal for base courses under flexible pavements."

    return {
        "soil_class_code": soil_code,
        "soil_class_name": soil_name,
        "description": description,
        "CBR_percent": round(cbr_est, 2),
        "group_index": round(gi_value, 2),
        "compaction_recommendation": compaction_rec,
        "classification_confidence": 95.0, # High confidence for rule-based matching
        "engine": "Geotechnical Rule Engine (USCS Fallback)"
    }


# =====================================================================
# 4. API Endpoints & Request Schema
# =====================================================================
class SoilGeotechMetrics(BaseModel):
    PL: float = Field(..., description="Plastic Limit (%)", ge=0.0, le=100.0)
    PI: float = Field(..., description="Plasticity Index (%)", ge=0.0, le=100.0)
    D10: float = Field(..., description="Diameter at which 10% of mass passes (mm)", gt=0.0)
    D30: float = Field(..., description="Diameter at which 30% of mass passes (mm)", gt=0.0)
    D60: float = Field(..., description="Diameter at which 60% of mass passes (mm)", gt=0.0)
    Cu: float = Field(None, description="Coefficient of Uniformity (calculated if omitted)")
    Cc: float = Field(None, description="Coefficient of Curvature (calculated if omitted)")
    OMC: float = Field(..., description="Optimum Moisture Content (%)", ge=0.0, le=100.0)
    MDD: float = Field(..., description="Maximum Dry Density (g/cm³)", ge=0.5, le=3.0)


# =====================================================================
# Continuous Learning Infrastructure Functions
# =====================================================================
def reload_calibrations_cache():
    """
    Loads calibrations.json from the root directory into the calibrations_cache.
    Handles missing file and malformed JSON errors gracefully.
    """
    global calibrations_cache
    calibrations_path = "calibrations.json"
    if os.path.exists(calibrations_path):
        try:
            with open(calibrations_path, 'r') as f:
                calibrations_cache = json.load(f)
            print(f" [SUCCESS] Calibration cache reloaded successfully: {calibrations_cache}")
            return True
        except Exception as e:
            print(f" [ERROR] Failed to load calibration file: {e}")
            calibrations_cache = {"error": f"Failed to parse calibration file: {e}"}
            return False
    else:
        print(" [INFO] calibrations.json does not exist. Calibration cache is empty.")
        calibrations_cache = {}
        return False


def reload_active_models():
    """
    Safely rebuilds the TensorFlow/Keras prediction session and reloads
    the Joblib scaler from the `./models/` directory.
    Dynamic imports are used to avoid startup crashes if libraries are missing.
    """
    global tf_model, joblib_scaler, model_status
    
    model_path = os.path.join("models", "soil_ai_cbr_regressor.h5")
    scaler_path = os.path.join("models", "soil_ai_preprocessor.pkl")
    
    # Check dependencies dynamically
    try:
        import tensorflow as tf
        import joblib
        has_libs = True
    except ImportError as e:
        print(f" [WARNING] Libraries required for cloud models (tensorflow/joblib) are missing: {e}")
        model_status["tf_model"] = f"Missing libraries: {e}"
        model_status["joblib_scaler"] = f"Missing libraries: {e}"
        has_libs = False
        
    if has_libs:
        # Load TensorFlow/Keras model
        if os.path.exists(model_path):
            try:
                tf_model = tf.keras.models.load_model(model_path)
                model_status["tf_model"] = f"Loaded Successfully ({os.path.basename(model_path)})"
                print(f" [LOAD OK] Keras model '{model_path}' loaded into memory.")
            except Exception as e:
                model_status["tf_model"] = f"Error loading model: {e}"
                print(f" [LOAD FAILED] Error loading Keras model: {e}")
        else:
            model_status["tf_model"] = "File Missing"
            print(f" [INFO] Keras model file not found: {model_path}")
            tf_model = None

        # Load Joblib scaler
        if os.path.exists(scaler_path):
            try:
                joblib_scaler = joblib.load(scaler_path)
                model_status["joblib_scaler"] = f"Loaded Successfully ({os.path.basename(scaler_path)})"
                print(f" [LOAD OK] Joblib scaler '{scaler_path}' loaded into memory.")
            except Exception as e:
                model_status["joblib_scaler"] = f"Error loading scaler: {e}"
                print(f" [LOAD FAILED] Error loading Joblib scaler: {e}")
        else:
            model_status["joblib_scaler"] = "File Missing"
            print(f" [INFO] Joblib scaler file not found: {scaler_path}")
            joblib_scaler = None
    else:
        tf_model = None
        joblib_scaler = None


def backup_current_models() -> str:
    """
    Creates a timestamped backup directory under `./models/backup/`
    and copies current active model files into it if they exist.
    Returns the path to the backup folder created, or None if no files were found to back up.
    """
    models_dir = "./models"
    backup_dir = os.path.join(models_dir, "backup")
    
    model_path = os.path.join(models_dir, "soil_ai_cbr_regressor.h5")
    scaler_path = os.path.join(models_dir, "soil_ai_preprocessor.pkl")
    
    model_exists = os.path.exists(model_path)
    scaler_exists = os.path.exists(scaler_path)
    
    if model_exists or scaler_exists:
        os.makedirs(backup_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        sub_backup_dir = os.path.join(backup_dir, f"backup_{timestamp}")
        os.makedirs(sub_backup_dir, exist_ok=True)
        
        if model_exists:
            shutil.copy2(model_path, os.path.join(sub_backup_dir, "soil_ai_cbr_regressor.h5"))
            print(f" [BACKUP] Copied model to {sub_backup_dir}")
        if scaler_exists:
            shutil.copy2(scaler_path, os.path.join(sub_backup_dir, "soil_ai_preprocessor.pkl"))
            print(f" [BACKUP] Copied preprocessor to {sub_backup_dir}")
            
        return sub_backup_dir
    return None


def count_input_items():
    """
    Counts files in `./data/inputs/` directory. Handles missing directory gracefully.
    """
    inputs_dir = os.path.join("data", "inputs")
    if not os.path.exists(inputs_dir):
        return 0
    try:
        items = os.listdir(inputs_dir)
        files = [f for f in items if os.path.isfile(os.path.join(inputs_dir, f))]
        return len(files)
    except Exception as e:
        print(f" [ERROR] Error counting files in inputs directory: {e}")
        return 0


def get_model_modified_date():
    """
    Returns the file modification ISO timestamp of `./models/soil_ai_cbr_regressor.h5`.
    """
    model_path = os.path.join("models", "soil_ai_cbr_regressor.h5")
    if not os.path.exists(model_path):
        return "N/A"
    try:
        mtime = os.path.getmtime(model_path)
        return datetime.fromtimestamp(mtime).isoformat()
    except Exception as e:
        print(f" [ERROR] Error getting model modified date: {e}")
        return "N/A"


@app.on_event("startup")
def startup_event():
    # Execute extraction
    extract_and_verify_archive()
    # Initialize ML models
    initialize_models()
    # Load calibrations cache
    reload_calibrations_cache()
    # Load active models
    reload_active_models()


@app.get("/")
def get_status(request: Request):
    accept = request.headers.get("accept", "")
    if "text/html" in accept:
        index_path = os.path.join("build", "web", "index.html")
        if os.path.exists(index_path):
            try:
                with open(index_path, "r", encoding="utf-8") as f:
                    return HTMLResponse(content=f.read(), status_code=200)
            except Exception as e:
                return HTMLResponse(content=f"<h3>Error reading Flutter web build: {e}</h3>", status_code=500)
        else:
            return HTMLResponse(
                content="<h3>Geotechnical Soil AI API Server is online. Flutter Web build not found.</h3>",
                status_code=200
            )
    return {
        "status": "online",
        "archive_extracted": os.path.exists(EXTRACT_DIR),
        "using_ml_models": using_ml_models,
        "loaded_models_log": model_status
    }


@app.post("/analyze-soil")
def analyze_soil(metrics: SoilGeotechMetrics):
    """
    Exposes POST /analyze-soil endpoint.
    Accepts raw engineering data and runs predictions via ML pipeline or fallback engine.
    """
    # 1. Compute Cu and Cc if they are omitted or zero
    d10 = metrics.D10
    d30 = metrics.D30
    d60 = metrics.D60
    
    cu = metrics.Cu
    cc = metrics.Cc
    
    if cu is None or cu <= 0.0:
        cu = d60 / d10 if d10 > 0 else 0.0
    if cc is None or cc <= 0.0:
        cc = (d30 * d30) / (d60 * d10) if (d60 * d10) > 0 else 0.0

    try:
        # 2. Check if we can run via Loaded ML Models
        if using_ml_models and "rf_classifier" in models and "cbr_regressor" in models:
            # Construct feature array matching feature_list.json exactly:
            # ["PL", "PI", "D10", "D30", "D60", "Cu", "Cc", "OMC", "MDD"]
            features = [
                metrics.PL,
                metrics.PI,
                d10,
                d30,
                d60,
                cu,
                cc,
                metrics.OMC,
                metrics.MDD
            ]
            
            # Predict Soil Classification Class ID
            clf = models["rf_classifier"]
            pred_class_id = int(clf.predict([features])[0])
            
            # Get class probabilities
            try:
                probs = clf.predict_proba([features])[0]
                confidence = float(probs[pred_class_id]) * 100
            except Exception:
                confidence = 85.0 # fallback confidence

            # Decode class ID using Label Encoder if available
            soil_code = "CL"
            if "label_encoder" in models:
                try:
                    soil_code = str(models["label_encoder"].inverse_transform([pred_class_id])[0])
                except Exception:
                    # fallback mappings
                    class_map = {0: "CL", 1: "CH", 2: "ML", 3: "MH", 4: "SW", 5: "SP"}
                    soil_code = class_map.get(pred_class_id, "CL")
            
            # Predict CBR value
            cbr_reg = models["cbr_regressor"]
            cbr_val = float(cbr_reg.predict([features])[0])
            cbr_val = max(1.0, min(100.0, cbr_val)) # realistic bounds
            
            # Predict Group Index
            gi_val = 0.0
            if "rf_gi_regressor" in models:
                gi_val = float(models["rf_gi_regressor"].predict([features])[0])
            else:
                # Calculate via formula if regressor is missing
                # Sieve #200 passing percentage proxy
                fines = 65.0 if d10 < 0.075 else 12.0
                LL = metrics.PL + metrics.PI
                f_term = max(0.0, fines - 35.0)
                ll_term = max(0.0, LL - 40.0)
                f_term_pi = max(0.0, fines - 15.0)
                pi_term = max(0.0, metrics.PI - 10.0)
                gi_val = (f_term * (0.2 + 0.005 * ll_term)) + (0.01 * f_term_pi * pi_term)
            
            gi_val = max(0.0, round(gi_val, 2))

            # Code descriptions & compaction recommendations
            descriptions = {
                "CL": "Lean Clay - low to medium plasticity inorganic clay. Moderate compressibility.",
                "CH": "Fat Clay - high plasticity inorganic clay. Expansive, low stability.",
                "ML": "Silt - low plasticity inorganic silt. Low cohesion, high frost risk.",
                "MH": "Elastic Silt - high plasticity inorganic silt. Compressible and elastic.",
                "SW": "Well-graded Sand - clean sand with balanced grain distribution. Outstanding load bearing.",
                "SP": "Poorly-graded Sand - clean, uniform sand. Low cohesion.",
                "GW": "Well-graded Gravel - clean, coarse gravelly mixtures. Excellent foundation material.",
                "GP": "Poorly-graded Gravel - clean, uniform gravelly soil. Highly permeable."
            }
            
            desc = descriptions.get(soil_code, "Geotechnical Soil Specimen")
            
            if cbr_val < 5.0:
                compaction_rec = "Very poor subgrade. Stabilization with lime or cement, or subgrade replacement/geogrid reinforcement is highly recommended."
            elif cbr_val < 10.0:
                compaction_rec = "Fair subgrade. Subgrade needs controlled moisture compaction. Good compaction at OMC (+/- 2%) is vital."
            elif cbr_val < 30.0:
                compaction_rec = "Good subgrade / base. Standard compaction to 95% MDD is sufficient. Suitable for sub-base courses."
            else:
                compaction_rec = "Excellent subgrade / base. Outstanding bearing capacity. Ideal for base courses under flexible pavements."

            return {
                "soil_class_code": soil_code,
                "soil_class_name": f"{soil_code} ({desc.split(' - ')[0]})",
                "description": desc,
                "CBR_percent": round(cbr_val, 2),
                "group_index": round(gi_val, 2),
                "compaction_recommendation": compaction_rec,
                "classification_confidence": round(confidence, 2),
                "engine": "Geotechnical ML Pipeline (Scikit-Learn Random Forest)"
            }
            
        else:
            # 3. Fallback engine
            return run_uscs_fallback_engine(
                metrics.PL, metrics.PI, d10, d30, d60, cu, cc, metrics.OMC, metrics.MDD
            )
            
    except Exception as e:
        print(f" [RUNTIME ERROR] Error during prediction pipeline: {e}")
        traceback.print_exc()
        # Fallback to the rule engine instead of crashing
        return run_uscs_fallback_engine(
            metrics.PL, metrics.PI, d10, d30, d60, cu, cc, metrics.OMC, metrics.MDD
        )
def calculate_log_interpolation(percent, diameters, passing):
    """
    Finds the particle diameter corresponding to a given percent passing
    using logarithmic interpolation.
    """
    pts = []
    for d, p in zip(diameters, passing):
        if d is not None and p is not None and d > 0:
            pts.append((d, p))
    pts.sort(key=lambda x: x[0], reverse=True)
    
    for i in range(len(pts) - 1):
        d1, p1 = pts[i]
        d2, p2 = pts[i+1]
        
        min_p = min(p1, p2)
        max_p = max(p1, p2)
        if min_p <= percent <= max_p:
            if abs(p2 - p1) < 1e-6:
                return d1
            log_d = math.log10(d1) + (percent - p1) / (p2 - p1) * (math.log10(d2) - math.log10(d1))
            return 10 ** log_d
            
    if len(pts) > 0:
        if percent < pts[-1][1]:
            d1, p1 = pts[-2] if len(pts) > 1 else (pts[-1][0] * 10, pts[-1][1] * 2)
            d2, p2 = pts[-1]
            if abs(p2 - p1) > 1e-6:
                log_d = math.log10(d1) + (percent - p1) / (p2 - p1) * (math.log10(d2) - math.log10(d1))
                return max(0.001, 10 ** log_d)
            return pts[-1][0]
        else:
            return pts[0][0]
    return 0.1

def classify_soil_uscs(gravel_pct, sand_pct, fines_pct, cu, cc):
    """
    Classifies the soil based on USCS grading criteria.
    """
    if fines_pct >= 50.0:
        if fines_pct > 70.0:
            return "Fat Clay"
        else:
            return "Lean Clay"
            
    if gravel_pct > sand_pct:
        # Gravel
        if fines_pct < 5.0:
            if cu >= 4.0 and 1.0 <= cc <= 3.0:
                name = "Well-Graded Gravel"
            else:
                name = "Poorly-Graded Gravel"
        elif fines_pct > 12.0:
            name = "Clayey Gravel" if fines_pct > 25.0 else "Silty Gravel"
        else:
            if cu >= 4.0 and 1.0 <= cc <= 3.0:
                name = "Well-Graded Gravel with Silt"
            else:
                name = "Poorly-Graded Gravel with Silt"
        if sand_pct >= 15.0:
            name += " with Sand"
        return name
    else:
        # Sand
        if fines_pct < 5.0:
            if cu >= 6.0 and 1.0 <= cc <= 3.0:
                name = "Well-Graded Sand"
            else:
                name = "Poorly-Graded Sand"
        elif fines_pct > 12.0:
            name = "Clayey Sand" if fines_pct > 25.0 else "Silty Sand"
        else:
            if cu >= 6.0 and 1.0 <= cc <= 3.0:
                name = "Well-Graded Sand with Silt"
            else:
                name = "Poorly-Graded Sand with Silt"
        if gravel_pct >= 15.0:
            name += " with Gravel"
        return name


class SieveSampleInput(BaseModel):
    sample_no: str = Field(..., description="Sample Number (e.g. BH-7)")
    depth: float = Field(..., description="Depth in meters")
    weight: float = Field(..., description="Weight of Sample")
    incremental_weights: List[float] = Field(..., description="13-row list of incremental weights in grams")


class ExportSieveRequest(BaseModel):
    region: str = Field(..., description="Region name")
    date_sampled: str = Field(..., description="Date Sampled string")
    tested_by: str = Field("", description="Tested By string")
    sample_description: str = Field(..., description="Sample Description")
    location: str = Field(..., description="Location of Sample")
    samples: List[SieveSampleInput] = Field(..., description="List of samples to inject as sheets")


@app.post("/api/export-sieve")
def export_sieve(req: ExportSieveRequest):
    """
    Cognitive Excel Export Engine endpoint.
    Loads the template sheet, inserts the logo, overwrites metadata cells,
    shifts row styles and writes incremental weights and formulas for each sample,
    calculates custom geotechnical parameters and USCS classification,
    and inserts signatures/stamp images before returning a multi-tab workbook.
    """
    try:
        template_path = "Excel_Table_Sample/Masika Master Template.xlsx"
        logo_path = "Excel_Table_Sample/Logo.png"
        stamp_path = "Excel_Table_Sample/stamp.png"
        engineer_signature_path = "Excel_Table_Sample/material engineer signature.png"
        technician_signature_path = "Excel_Table_Sample/Geotechnical Laboratory Technician Signature.png"
        
        if not os.path.exists(template_path):
            raise HTTPException(status_code=404, detail="Template workbook not found.")
            
        wb = openpyxl.load_workbook(template_path)
        original_ws = wb.active
        
        # Helper function to copy cell style
        def copy_style(src, dest):
            if src.has_style:
                dest.font = copy(src.font)
                dest.border = copy(src.border)
                dest.fill = copy(src.fill)
                dest.alignment = copy(src.alignment)
                dest.number_format = copy(src.number_format)
                
        # Loop over each sample in the request and build its worksheet
        for sample in req.samples:
            ws = wb.copy_worksheet(original_ws)
            ws.title = f"{sample.sample_no} ({sample.depth:.1f}m)"
            
            # 1. Header Metadata Form Injection Map
            ws['B7'] = f"Region: {req.region}"
            ws['B11'] = f"Date Sampled: {req.date_sampled}"
            
            # If Tested By is empty, leave Cell G12 completely blank for post-download manual entries
            if req.tested_by and req.tested_by.strip():
                ws['G12'] = req.tested_by
            else:
                ws['G12'] = ""
                
            ws['E9'] = sample.sample_no
            ws['E10'] = sample.depth
            ws['E13'] = req.location
            ws['E14'] = sample.weight
            
            # Labels and sieve dia for Column B and C (rows 19 to 31)
            sieve_labels = [
                ('3"', 75.0),
                ('1-1/2"', 37.5),
                ('3/4"', 19.0),
                ('0.375', 9.5),
                ('No. 4', 4.75),
                ('No. 10', 2.36),
                ('No. 16', 1.18),
                ('No. 30', 0.60),
                ('No. 40', 0.425),
                ('No. 50', 0.30),
                ('No. 100', 0.150),
                ('No. 200', 0.075),
                ('Passing 0.075', None)
            ]
            
            for idx, (name, size) in enumerate(sieve_labels):
                r = 19 + idx
                ws[f"B{r}"] = name
                if size is not None:
                    ws[f"C{r}"] = size
                    
            # Mass Data Matrix Overwrite (Column D)
            for idx, val in enumerate(sample.incremental_weights):
                if idx < len(sieve_labels):
                    r = 19 + idx
                    ws[f"D{r}"] = val
                    
            # Set up formulas in Columns E, F, G for Rows 19-31
            ws['E19'] = "=D19"
            ws['F19'] = "=(D19/$E$14)*100"
            ws['G19'] = "=100-F19"
            
            for r in range(20, 32):
                ws[f"E{r}"] = f"=E{r-1}+D{r}"
                ws[f"F{r}"] = f"=(D{r}/$E$14)*100"
                ws[f"G{r}"] = f"=G{r-1}-F{r}"
                
            # Set Upper/Lower Limit headers at row 17 (matching % Passing style in G17)
            ws['H17'] = "Upper Limit"
            ws['I17'] = "Lower Limits"
            copy_style(ws['G17'], ws['H17'])
            copy_style(ws['G17'], ws['I17'])
            
            # Totals Row (Row 32)
            ws['B32'] = "Totals:"
            ws['B32'].font = openpyxl.styles.Font(name=ws['B32'].font.name, size=ws['B32'].font.size, bold=True, color=ws['B32'].font.color)
            ws['C32'] = "MM"
            ws['D32'] = "=SUM(D19:D31)"
            ws['E32'] = "=E31"
            ws['F32'] = "=SUM(F19:F31)"
            ws['G32'] = None
            
            # Merge B31:C31 for Passing 0.075 and B32:C32 for Totals
            try:
                ws.merge_cells('B31:C31')
            except ValueError:
                pass
            try:
                ws.merge_cells('B32:C32')
            except ValueError:
                pass

            # Calculate cumulative passing percentages in Python for interpolation and classification
            total_wt = sample.weight if sample.weight > 0 else 500.0
            cum_wt = 0.0
            passing_pcts = []
            for w in sample.incremental_weights:
                cum_wt += w
                pct_retained = (cum_wt / total_wt) * 100.0
                pct_passing = 100.0 - pct_retained
                passing_pcts.append(pct_passing)

            # Diameters for the first 12 sieves
            diameters = [75.0, 37.5, 19.0, 9.5, 4.75, 2.36, 1.18, 0.60, 0.425, 0.30, 0.150, 0.075]

            # Logarithmic interpolation for D60, D30, D10
            d60 = calculate_log_interpolation(60.0, diameters, passing_pcts[:12])
            d30 = calculate_log_interpolation(30.0, diameters, passing_pcts[:12])
            d10 = calculate_log_interpolation(10.0, diameters, passing_pcts[:12])

            cu = d60 / d10 if d10 > 0 else 0.0
            cc = (d30 * d30) / (d60 * d10) if (d60 * d10) > 0 else 0.0

            # Calculate fractions
            gravel_val = passing_pcts[0] - passing_pcts[5]
            sand_val = passing_pcts[5] - passing_pcts[11]
            fines_val = passing_pcts[11]

            # Soil Classification
            soil_class = classify_soil_uscs(gravel_val, sand_val, fines_val, cu, cc)
            ws['G14'] = soil_class
            # 3. Summary Cards and Coefficients (Rows 62-64)
            font_arial = openpyxl.styles.Font(name="Arial", size=12)
            medium_side = openpyxl.styles.Side(border_style="medium", color="000000")
            thin_side = openpyxl.styles.Side(border_style="thin", color="000000")
            
            # Summary Labels, Formulas and Unit in B62:F64
            summary_labels = ["Gravel", "Sand", "Fine"]
            summary_formulas = ["=G19-G24", "=G24-G30", "=G30"]
            
            for idx, (label, formula) in enumerate(zip(summary_labels, summary_formulas)):
                r_card = 62 + idx
                
                # Merge B and C for label
                try:
                    ws.merge_cells(start_row=r_card, start_column=2, end_row=r_card, end_column=3)
                except ValueError:
                    pass
                ws[f'B{r_card}'] = label
                ws[f'B{r_card}'].font = font_arial
                ws[f'B{r_card}'].alignment = openpyxl.styles.Alignment(horizontal="left", vertical="center")
                
                # Merge D and E for value
                try:
                    ws.merge_cells(start_row=r_card, start_column=4, end_row=r_card, end_column=5)
                except ValueError:
                    pass
                ws[f'D{r_card}'] = formula
                ws[f'D{r_card}'].font = font_arial
                ws[f'D{r_card}'].alignment = openpyxl.styles.Alignment(horizontal="center", vertical="center")
                ws[f'D{r_card}'].number_format = '0.00'
                
                # F is for %
                ws[f'F{r_card}'] = "%"
                ws[f'F{r_card}'].font = font_arial
                ws[f'F{r_card}'].alignment = openpyxl.styles.Alignment(horizontal="center", vertical="center")
            
            # G62 and G64 for CU and CC values
            ws['G62'] = f"CU (D60/D10): {round(cu, 1)}"
            ws['G62'].font = font_arial
            ws['G62'].alignment = openpyxl.styles.Alignment(horizontal="left", vertical="center")
            
            ws['G63'] = None
            
            ws['G64'] = f"CC (D30^2/(D60*D10)):{round(cc, 1)}"
            ws['G64'].font = font_arial
            ws['G64'].alignment = openpyxl.styles.Alignment(horizontal="left", vertical="center")

            # Set exact borders for Rows 62-64 (columns B to G)
            for r in range(62, 65):
                for c in range(2, 8):
                    cell = ws.cell(row=r, column=c)
                    
                    left_b = None
                    right_b = None
                    top_b = None
                    bottom_b = None
                    
                    # Left edge
                    if c == 2:
                        left_b = medium_side
                    elif c in [4, 6]:
                        left_b = thin_side
                        
                    # Right edge
                    if c == 7:
                        right_b = medium_side
                    elif c in [3, 5, 6]:
                        right_b = thin_side
                        
                    # Top edge
                    if r == 62:
                        top_b = medium_side
                    else:
                        if c == 7:
                            top_b = None
                        else:
                            top_b = thin_side
                        
                    # Bottom edge
                    if r == 64:
                        if c == 7:
                            bottom_b = None
                        else:
                            bottom_b = medium_side
                    else:
                        if c == 7:
                            bottom_b = None
                        else:
                            bottom_b = thin_side
                            
                    cell.border = openpyxl.styles.Border(left=left_b, right=right_b, top=top_b, bottom=bottom_b)

            # 4. Draw Signature Blocks at B65:D67 (Left) and E65:G67 (Right)
            ws['B65'] = f"Contractor (Material Engineer): {req.tested_by or 'Sylvester'}"
            ws['B65'].font = font_arial
            ws['E65'] = "Geotechnical Laboratory Technician: Sallieu"
            ws['E65'].font = font_arial
            
            ws['B66'] = f"Date:{req.date_sampled}"
            ws['B66'].font = font_arial
            ws['E66'] = f"Date:{req.date_sampled}"
            ws['E66'].font = font_arial
            
            ws['B67'] = "Signature:"
            ws['B67'].font = font_arial
            ws['E67'] = "Signature:"
            ws['E67'].font = font_arial

            for r in range(65, 68):
                for c in range(2, 8):
                    cell = ws.cell(row=r, column=c)
                    
                    left_b = None
                    right_b = None
                    top_b = None
                    bottom_b = None
                    
                    # Left signature box B:D
                    if c in [2, 3, 4]:
                        if c == 2:
                            left_b = medium_side
                        elif c == 3:
                            left_b = thin_side
                        
                        if c == 4:
                            right_b = medium_side
                            
                        if r == 65:
                            top_b = medium_side
                        if r == 67:
                            bottom_b = medium_side
                            
                    # Right signature box E:G
                    if c in [5, 6, 7]:
                        if c == 5:
                            left_b = medium_side
                        if c == 7:
                            right_b = medium_side
                            
                        if r == 65:
                            top_b = medium_side
                        if r == 67:
                            bottom_b = medium_side
                            
                    cell.border = openpyxl.styles.Border(left=left_b, right=right_b, top=top_b, bottom=bottom_b)

            # 5. Graphical Asset Injections (Logo, Stamp, Signatures with exact scaling)
            if os.path.exists(logo_path):
                img_logo = OpenpyxlImage(logo_path)
                ws.add_image(img_logo, 'B2')
                
            if os.path.exists(engineer_signature_path):
                img_eng = OpenpyxlImage(engineer_signature_path)
                orig_w, orig_h = img_eng.width, img_eng.height
                img_eng.height = 45
                img_eng.width = int(orig_w * (45 / orig_h))
                ws.add_image(img_eng, 'C67')
                
            if os.path.exists(technician_signature_path):
                img_tech = OpenpyxlImage(technician_signature_path)
                orig_w_t, orig_h_t = img_tech.width, img_tech.height
                img_tech.height = 45
                img_tech.width = int(orig_w_t * (45 / orig_h_t))
                ws.add_image(img_tech, 'F67')
                
            if os.path.exists(stamp_path):
                img_stamp = OpenpyxlImage(stamp_path)
                img_stamp.width = 95
                img_stamp.height = 95
                ws.add_image(img_stamp, 'C65')

            # 6. Initialize Semi-Log Scatter Chart (B34 to H61)
            chart = ScatterChart()
            chart.title = "PARTICLE SIZE DISTRIBUTION CURVE"
            
            # Explicitly prevent axes from being deleted/hidden to ensure numbers and tick labels show
            chart.x_axis.delete = False
            chart.y_axis.delete = False
            
            # Format title font (centered, bold, black, 12pt Arial)
            cp_title = CharacterProperties(latin=DrawFont(typeface='Arial'), sz=1200, b=True, solidFill=ColorChoice(srgbClr="000000"))
            pp_title = ParagraphProperties(defRPr=cp_title)
            chart.title.txPr = RichText(p=[Paragraph(pPr=pp_title, endParaRPr=cp_title)])
            
            # TwoCellAnchor for placing chart directly UNDER the table, spanning Columns B to G (B34 to H61)
            chart.anchor = TwoCellAnchor()
            chart.anchor._from.col = 1   # Column B
            chart.anchor._from.row = 33  # Row 34 (0-indexed)
            chart.anchor.to.col = 7      # Column H (so Column G is included)
            chart.anchor.to.row = 60     # Row 61 (so Row 60 is included)

            # Chart Area background and border (solid white, no border)
            chart.graphical_properties = GraphicalProperties(solidFill="FFFFFF")
            chart.graphical_properties.line.noFill = True

            # Plot Area background and border (solid white, black border)
            chart.plot_area.graphicalProperties = GraphicalProperties(solidFill="FFFFFF")
            chart.plot_area.graphicalProperties.line.solidFill = "000000"
            chart.plot_area.graphicalProperties.line.width = 12700 # 1.0pt

            # X-axis configuration (base-10 log scale from 0.01 to 100)
            chart.x_axis.title = "GRAIN SIZE IN MILLIMETERS (mm)"
            chart.x_axis.scaling.logBase = 10
            chart.x_axis.scaling.min = 0.01
            chart.x_axis.scaling.max = 100
            chart.x_axis.scaling.orientation = "minMax"
            chart.x_axis.crosses = "min"

            # Major gridlines (dark black, 1.0pt)
            chart.x_axis.majorGridlines = ChartLines()
            chart.x_axis.majorGridlines.spPr = GraphicalProperties(ln=LineProperties(solidFill="000000", w=12700))

            # Minor gridlines (dark gray, 0.5pt)
            chart.x_axis.minorGridlines = ChartLines()
            chart.x_axis.minorGridlines.spPr = GraphicalProperties(ln=LineProperties(solidFill="555555", w=6350))

            # Y-axis configuration (linear, 0.00 to 120.00, major units 10)
            chart.y_axis.title = "PERCENTAGE PASSING (%)"
            chart.y_axis.scaling.min = 0.00
            chart.y_axis.scaling.max = 120.00
            chart.y_axis.majorUnit = 10
            chart.y_axis.numFmt = "0.00"
            chart.y_axis.crosses = "min"

            # Major gridlines (black, 0.5pt)
            chart.y_axis.majorGridlines = ChartLines()
            chart.y_axis.majorGridlines.spPr = GraphicalProperties(ln=LineProperties(solidFill="000000", w=6350))

            # Apply manual layout coordinates to prevent overlapping of title, legend, and axes labels
            chart.layout = Layout(
                manualLayout=ManualLayout(
                    layoutTarget="inner",
                    xMode="edge",
                    yMode="edge",
                    wMode="factor",
                    hMode="factor",
                    x=0.119745222929936,
                    y=0.190231362467866,
                    w=0.742675159235669,
                    h=0.616966580976863
                )
            )
            
            chart.title.layout = Layout(
                manualLayout=ManualLayout(
                    xMode="edge",
                    yMode="edge",
                    wMode="factor",
                    hMode="factor",
                    x=0.303933917661863,
                    y=0.0334193363671607
                )
            )

            chart.x_axis.title.layout = Layout(
                manualLayout=ManualLayout(
                    xMode="edge",
                    yMode="edge",
                    wMode="factor",
                    hMode="factor",
                    x=0.373248261241546,
                    y=0.876606609352457
                )
            )

            chart.y_axis.title.layout = Layout(
                manualLayout=ManualLayout(
                    xMode="edge",
                    yMode="edge",
                    wMode="factor",
                    hMode="factor",
                    x=0.0191082127319499,
                    y=0.31362462654672
                )
            )

            # Series definitions (rows 19 to 30)
            xvalues = Reference(ws, min_col=3, min_row=19, max_row=30)
            yvalues1 = Reference(ws, min_col=7, min_row=19, max_row=30)
            yvalues2 = Reference(ws, min_col=8, min_row=19, max_row=30)
            yvalues3 = Reference(ws, min_col=9, min_row=19, max_row=30)

            s1 = Series(yvalues1, xvalues, title="GRAVEL - SAND")
            s2 = Series(yvalues2, xvalues, title="Upper Limit")
            s3 = Series(yvalues3, xvalues, title="Lower Limits")

            # Series 1: Orange, square markers
            s1.graphicalProperties.line.solidFill = "F26522"
            s1.graphicalProperties.line.width = 25400 # 2.0pt
            s1.marker.symbol = "square"
            s1.marker.size = 5
            s1.marker.graphicalProperties.solidFill = "F26522"
            s1.marker.graphicalProperties.line.solidFill = "F26522"

            # Series 2: Black dashed, no markers
            s2.graphicalProperties.line.solidFill = "000000"
            s2.graphicalProperties.line.prstDash = "dash"
            s2.graphicalProperties.line.width = 12700 # 1.0pt
            s2.marker.symbol = "none"

            # Series 3: Black dashed, triangle markers
            s3.graphicalProperties.line.solidFill = "000000"
            s3.graphicalProperties.line.prstDash = "dash"
            s3.graphicalProperties.line.width = 12700 # 1.0pt
            s3.marker.symbol = "triangle"
            s3.marker.size = 5
            s3.marker.graphicalProperties.solidFill = "000000"
            s3.marker.graphicalProperties.line.solidFill = "000000"

            chart.series.append(s1)
            chart.series.append(s2)
            chart.series.append(s3)

            # Legend position: bottom center
            chart.legend.position = "b"

            ws.add_chart(chart)



        # Remove the original unpopulated worksheet template from workbook
        wb.remove(original_ws)
        
        # Save workbook to memory buffer
        file_stream = io.BytesIO()
        wb.save(file_stream)
        file_stream.seek(0)
        
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename=Sieve_Analysis_Report.xlsx"}
        )
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate Excel report: {str(e)}")


# =====================================================================
# 5. Continuous Learning API Endpoints
# =====================================================================
@app.post("/api/admin/train-local")
def train_local():
    """
    POST endpoint to programmatically execute the train_light.py script.
    Captures stdout/stderr. If successful (exit code 0), reloads calibrations.
    """
    script_path = "train_light.py"
    if not os.path.exists(script_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Calibration training script '{script_path}' was not found in the root directory."
        )
        
    try:
        # Run using sys.executable to ensure we use the same Python interpreter environment
        result = subprocess.run(
            [sys.executable, script_path],
            capture_output=True,
            text=True,
            check=False
        )
        
        if result.returncode == 0:
            # Trigger reload function
            reload_calibrations_cache()
            return {
                "status": "success",
                "message": "Local retraining completed successfully and calibrations reloaded.",
                "exit_code": result.returncode,
                "stdout": result.stdout,
                "stderr": result.stderr,
                "calibrations": calibrations_cache
            }
        else:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail={
                    "status": "failed",
                    "message": f"Retraining script exited with non-zero code {result.returncode}",
                    "exit_code": result.returncode,
                    "stdout": result.stdout,
                    "stderr": result.stderr
                }
            )
    except Exception as e:
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={
                "status": "failed",
                "message": f"Exception occurred while launching retraining script: {str(e)}",
                "error": traceback.format_exc()
            }
        )


@app.post("/api/admin/train-light-model")
def train_light_model():
    """
    POST endpoint to programmatically execute local training and calibrations update.
    Reuses the train_local routine.
    """
    return train_local()


@app.post("/api/admin/sync-cloud-model")
def sync_cloud_model():
    """
    POST endpoint to sync production models from local Google Drive desktop directory.
    If new models exist, backs up current models, overwrites active ones, and reloads them.
    """
    gdrive_path = "C:/Users/pc/Google Drive/MyDrive/Soil_AI_Dataset/production_models/"
    gdrive_model_path = os.path.join(gdrive_path, "soil_ai_cbr_regressor.h5")
    gdrive_scaler_path = os.path.join(gdrive_path, "soil_ai_preprocessor.pkl")
    
    if not os.path.exists(gdrive_path):
        return {
            "status": "ignored",
            "message": f"Google Drive directory '{gdrive_path}' is not accessible on this system."
        }
        
    model_exists = os.path.exists(gdrive_model_path)
    scaler_exists = os.path.exists(gdrive_scaler_path)
    
    if not model_exists and not scaler_exists:
        return {
            "status": "ignored",
            "message": "No new model or preprocessor files found in the Google Drive production directory."
        }
        
    models_dir = "./models"
    os.makedirs(models_dir, exist_ok=True)
    
    # 1. Create a timestamped backup of current active models
    backup_folder = backup_current_models()
    
    dest_model_path = os.path.join(models_dir, "soil_ai_cbr_regressor.h5")
    dest_scaler_path = os.path.join(models_dir, "soil_ai_preprocessor.pkl")
    
    files_copied = []
    
    try:
        # 2. Overwrite production models
        if model_exists:
            shutil.copy2(gdrive_model_path, dest_model_path)
            files_copied.append("soil_ai_cbr_regressor.h5")
        if scaler_exists:
            shutil.copy2(gdrive_scaler_path, dest_scaler_path)
            files_copied.append("soil_ai_preprocessor.pkl")
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to copy model files from Google Drive: {e}"
        )
        
    # 3. Reload active models on the fly
    reload_active_models()
    
    return {
        "status": "success",
        "message": f"Successfully synced and hot-swapped {len(files_copied)} model files from Google Drive.",
        "files_synced": files_copied,
        "backup_created": backup_folder,
        "model_status": {
            "tf_model": model_status.get("tf_model"),
            "joblib_scaler": model_status.get("joblib_scaler")
        }
    }


@app.post("/api/admin/deploy-model")
async def deploy_model(
    model_file: UploadFile = File(..., description="The new soil_ai_cbr_regressor.h5 file"),
    preprocessor_file: UploadFile = File(..., description="The new soil_ai_preprocessor.pkl file")
):
    """
    POST endpoint to deploy a new TensorFlow model and Joblib scaler preprocessor.
    Creates a backup of existing files, overwrites them with the uploaded ones,
    and hot-swaps them in the active runtime memory.
    """
    models_dir = "./models"
    os.makedirs(models_dir, exist_ok=True)
    
    # 1. Create a timestamped backup of current files
    backup_folder = backup_current_models()
    
    model_path = os.path.join(models_dir, "soil_ai_cbr_regressor.h5")
    scaler_path = os.path.join(models_dir, "soil_ai_preprocessor.pkl")
    
    # 2. Overwrite files
    try:
        with open(model_path, "wb") as buffer:
            shutil.copyfileobj(model_file.file, buffer)
            
        with open(scaler_path, "wb") as buffer:
            shutil.copyfileobj(preprocessor_file.file, buffer)
            
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to write uploaded model binaries: {e}"
        )
        
    # 3. Reload active models on the fly
    reload_active_models()
    
    return {
        "status": "success",
        "message": "Models deployed and hot-swapped successfully.",
        "backup_created": backup_folder,
        "model_status": {
            "tf_model": model_status.get("tf_model"),
            "joblib_scaler": model_status.get("joblib_scaler")
        }
    }


@app.get("/api/system-status")
def get_system_status():
    """
    GET endpoint returning global server and infrastructure status,
    including model modification times and pending input files count.
    """
    h5_modified_date = get_model_modified_date()
    parsed_inputs_count = count_input_items()
    
    return {
        "status": "online",
        "archive_extracted": os.path.exists(EXTRACT_DIR),
        "using_ml_models": using_ml_models,
        "loaded_models_log": model_status,
        "active_h5_modified_date": h5_modified_date,
        "parsed_inputs_count": parsed_inputs_count,
        "calibrations": calibrations_cache
    }


@app.post("/api/parse-sieve-image")
def parse_sieve_image(file: UploadFile = File(...)):
    """
    Parses an uploaded sieve log sheet image or PDF document using Gemini 1.5 Flash to extract:
    - borehole_id
    - location
    - date
    - samples (depth, weight, incremental sieve weights)
    """
    import base64
    import urllib.request
    import urllib.error
    
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key or api_key.strip() == "" or api_key == "YOUR_GEMINI_API_KEY_HERE":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="GEMINI_API_KEY is not set or has the placeholder value in the .env file. Please open the .env file in the project folder, paste your valid Gemini API Key, and restart the server."
        )

    try:
        # Read file contents
        contents = file.file.read()
        
        # Base64 encode the file
        file_base64 = base64.b64encode(contents).decode("utf-8")
        
        # Deduce mime type
        mime_type = file.content_type
        if not mime_type or mime_type == "application/octet-stream":
            filename = file.filename.lower() if file.filename else ""
            if filename.endswith(".pdf"):
                mime_type = "application/pdf"
            elif filename.endswith(".jpg") or filename.endswith(".jpeg"):
                mime_type = "image/jpeg"
            elif filename.endswith(".png"):
                mime_type = "image/png"
            else:
                mime_type = "image/png"  # fallback

        # Prepare the request to Gemini API
        gemini_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"
        
        prompt = """
        You are an expert Geotechnical Engineering AI assistant.
        Analyze the uploaded laboratory sieve analysis log sheet image or PDF document and extract the sample metadata and sieve data.
        
        Identify:
        1. Borehole ID (e.g. "BH-7" or "BH-01")
        2. Location / Project Site (e.g. "Masika" or "Sector 4")
        3. Date (e.g. "13/05/26" or "14/05/26")
        4. Sieve weights retained in grams for each sample column.
           Note that sieve rows in the document typically map to:
           - 75mm or 75.0mm (map to "75.0")
           - 37.5mm (map to "37.5")
           - 19.0mm (map to "19.0")
           - 9.5mm (map to "9.5")
           - 4.75mm (map to "4.75")
           - 2.36mm (map to "2.36")
           - 1.18mm (map to "1.18")
           - 0.60mm or 0.600mm (map to "0.60")
           - 0.425mm (map to "0.425")
           - 0.30mm or 0.300mm (map to "0.30")
           - 0.150mm (map to "0.150")
           - 0.075mm (map to "0.075")
           - Pan or Passing 0.075 (map to "0.001")
        
        For each sample column:
        - Determine its depth (m) from the column header (e.g. "20.0m", "6.0m", "5.0m", "4.0m", "12.0m", "10.0m", "3.0m").
        - Determine its initial total dry weight (g) from the "Wt=" or "Weight=" or "wt=" field in the column header.
        - Extract the incremental weights retained on each sieve size. Double numbers like "00" or "-" or blank entries mean 0.0.
        
        You must return a raw JSON object matching the following structure:
        {
          "borehole_id": "BH-7",
          "location": "Masika",
          "date": "13/05/26",
          "samples": [
            {
              "depth": 20.0,
              "weight": 259.5,
              "sieve_weights": {
                "75.0": 0.0,
                "37.5": 0.0,
                "19.0": 0.0,
                "9.5": 0.0,
                "4.75": 0.0,
                "2.36": 0.0,
                "1.18": 2.0,
                "0.60": 5.4,
                "0.425": 12.0,
                "0.30": 55.1,
                "0.150": 159.3,
                "0.075": 24.4,
                "0.001": 1.3
              }
            }
          ]
        }
        
        Make sure the returned JSON contains all the samples in the document. Return ONLY valid JSON, do not wrap in markdown blocks.
        """
        
        payload = {
            "contents": [
                {
                    "parts": [
                        {"text": prompt},
                        {
                            "inlineData": {
                                "mimeType": mime_type,
                                "data": file_base64
                            }
                        }
                    ]
                }
            ],
            "generationConfig": {
                "responseMimeType": "application/json"
            }
        }
        
        req_data = json.dumps(payload).encode("utf-8")
        req = urllib.request.Request(
            gemini_url,
            data=req_data,
            headers={"Content-Type": "application/json"},
            method="POST"
        )
        
        with urllib.request.urlopen(req) as response:
            res_body = response.read().decode("utf-8")
            res_json = json.loads(res_body)
            
            # Extract response text
            extracted_text = res_json["candidates"][0]["content"]["parts"][0]["text"]
            
            # Parse the extracted text as JSON
            parsed_data = json.loads(extracted_text.strip())
            return parsed_data
            
    except urllib.error.HTTPError as e:
        err_body = e.read().decode("utf-8") if e else ""
        print(f"Gemini API HTTP Error: {e.code} - {err_body}")
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"Gemini API returned an error: {err_body or str(e)}"
        )
    except Exception as e:
        print(f"Error calling Gemini Vision API: {e}")
        traceback.print_exc()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to process image: {str(e)}"
        )


# =====================================================================
# Word Document Converter Engine (Philip's Converter: Ultra Edition)
# =====================================================================
def set_cell_background(cell, hex_color: str):
    """
    Sets the background shading color of a Word table cell.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def get_clean_hex_color(cell):
    """
    Reads the background fill color from an openpyxl cell and returns it
    as a clean 6-digit RGB hex string. Returns None if default/empty/white/black.
    """
    if not cell.fill or not cell.fill.start_color:
        return None
        
    color_str = str(cell.fill.start_color.rgb)
    if not color_str or not re.match(r'^[0-9a-fA-F]+$', color_str):
        return None
        
    if len(color_str) == 8:
        rgb_color = color_str[2:]
        if rgb_color.upper() in ["FFFFFF", "000000"]:
            return None
        return rgb_color
    elif len(color_str) == 6:
        if color_str.upper() in ["FFFFFF", "000000"]:
            return None
        return color_str
        
    return None

def generate_matplotlib_chart(df: pd.DataFrame, sheet_name: str) -> io.BytesIO:
    """
    Generates a cloud-safe bar or line chart using Matplotlib based on Dataframe values.
    Returns a BytesIO binary stream containing the PNG.
    """
    num_cols = df.select_dtypes(include=['number']).columns.tolist()
    cat_cols = df.select_dtypes(exclude=['number']).columns.tolist()
    
    if not num_cols:
        return None
        
    y_col = num_cols[0]
    x_col = cat_cols[0] if cat_cols else None
    
    fig, ax = plt.subplots(figsize=(6.5, 3.8))
    try:
        if x_col:
            ax.bar(df[x_col].astype(str), df[y_col], color="#4F46E5", width=0.45, edgecolor="#3730A3", linewidth=1)
            ax.set_xlabel(str(x_col), fontsize=9, fontweight='bold', fontname='Arial')
        else:
            ax.plot(df[y_col], marker='o', color="#4F46E5", linewidth=2.5, markersize=6)
            ax.set_xlabel("Index", fontsize=9, fontweight='bold', fontname='Arial')
            
        ax.set_ylabel(str(y_col), fontsize=9, fontweight='bold', fontname='Arial')
        ax.set_title(f"{sheet_name} - Visual Metrics Representation", fontsize=11, fontweight='bold', fontname='Arial', color="#1E1B4B")
        
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['left'].set_color('#E2E8F0')
        ax.spines['bottom'].set_color('#E2E8F0')
        ax.grid(axis='y', linestyle=':', color="#CBD5E1", alpha=0.7)
        
        plt.xticks(rotation=15, fontsize=8, fontname='Arial')
        plt.yticks(fontsize=8, fontname='Arial')
        plt.tight_layout()
        
        img_buf = io.BytesIO()
        plt.savefig(img_buf, format='png', dpi=150, bbox_inches='tight')
        img_buf.seek(0)
        return img_buf
    except Exception as e:
        print(f"[!] Warning: Matplotlib chart drawing failed in web endpoint. Detail: {e}")
        return None
    finally:
        plt.close(fig)

def convert_excel_to_docx_bytes(excel_bytes: bytes) -> io.BytesIO:
    """
    Performs the full high-fidelity document conversion pipeline entirely in memory (RAM).
    """
    try:
        wb_openpyxl = openpyxl.load_workbook(io.BytesIO(excel_bytes), data_only=True)
    except Exception as e:
        raise ValueError(f"Invalid Excel Workbook file. Technical detail: {e}")
        
    sheet_names = wb_openpyxl.sheetnames
    if not sheet_names:
        raise ValueError("The workbook contains no active worksheets.")
        
    doc = docx.Document()
    
    # Page Margin constraints
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Centered Elegant Cover Page
    for _ in range(8):
        doc.add_paragraph()
        
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("PHILIP'S CONVERTER")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.name = 'Arial'
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run(
        "Advanced Enterprise Executive Report Generation\n"
        "Tables, Auto-Formatting, and Chart Integration Engine"
    )
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.italic = True
    subtitle_run.font.name = 'Arial'
    
    doc.add_page_break()
    
    total_sheets = len(sheet_names)
    
    for index, sheet_name in enumerate(sheet_names):
        ws = wb_openpyxl[sheet_name]
        
        header = doc.add_heading(level=1)
        header.paragraph_format.space_before = Pt(12)
        header.paragraph_format.space_after = Pt(12)
        h_run = header.add_run(f"Section: {sheet_name}")
        h_run.font.name = 'Arial'
        h_run.font.size = Pt(18)
        h_run.bold = True
        
        # Scan for merged ranges
        merged_map = {}
        for r_range in ws.merged_cells.ranges:
            min_row, min_col, max_row_range, max_col_range = r_range.min_row, r_range.min_col, r_range.max_row, r_range.max_col
            for r in range(min_row, max_row_range + 1):
                for c in range(min_col, max_col_range + 1):
                    merged_map[(r, c)] = (min_row, min_col, max_row_range, max_col_range)
                    
        # Find tight bounding box
        max_row = 0
        max_col = 0
        for r in range(1, ws.max_row + 1):
            for c in range(1, ws.max_column + 1):
                cell = ws.cell(row=r, column=c)
                has_color = get_clean_hex_color(cell) is not None
                if cell.value is not None or has_color or (r, c) in merged_map:
                    max_row = max(max_row, r)
                    max_col = max(max_col, c)
                    
        if max_row == 0 or max_col == 0:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run("[No active data table found on this worksheet]")
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.italic = True
        else:
            table = doc.add_table(rows=max_row, cols=max_col)
            table.autofit = False
            
            # Column auto-fitting widths calculations
            col_widths_inches = []
            for col_idx in range(max_col):
                max_char_len = 5
                for row_idx in range(max_row):
                    if (row_idx + 1, col_idx + 1) in merged_map:
                        min_row, min_col, _, _ = merged_map[(row_idx + 1, col_idx + 1)]
                        if (row_idx + 1) != min_row or (col_idx + 1) != min_col:
                            continue
                            
                    cell = ws.cell(row=row_idx + 1, column=col_idx + 1)
                    val_str = str(cell.value) if cell.value is not None else ""
                    max_char_len = max(max_char_len, len(val_str))
                
                calculated_width = max_char_len * 0.1
                final_width = max(0.8, min(3.5, calculated_width))
                col_widths_inches.append(final_width)
                
            try:
                table.style = 'Table Grid'
            except Exception:
                pass
                
            merged_already = set()
            warning_indicators = ["fail", "incomplete", "overdue", "critical"]
            
            for r in range(1, max_row + 1):
                # Row Alert scanning
                has_alert = False
                for c in range(1, max_col + 1):
                    cell = ws.cell(row=r, column=c)
                    val_str = str(cell.value).lower() if cell.value is not None else ""
                    if any(indicator in val_str for indicator in warning_indicators):
                        has_alert = True
                        break
                        
                for c in range(1, max_col + 1):
                    cell = ws.cell(row=r, column=c)
                    word_cell = table.cell(r - 1, c - 1)
                    
                    is_merged = (r, c) in merged_map
                    is_top_left = True
                    
                    if is_merged:
                        min_row, min_col, max_row_range, max_col_range = merged_map[(r, c)]
                        if r != min_row or c != min_col:
                            is_top_left = False
                        else:
                            range_key = (min_row, min_col, max_row_range, max_col_range)
                            if range_key not in merged_already:
                                lim_max_row = min(max_row_range, max_row)
                                lim_max_col = min(max_col_range, max_col)
                                if lim_max_row > min_row or lim_max_col > min_col:
                                    word_cell.merge(table.cell(lim_max_row - 1, lim_max_col - 1))
                                merged_already.add(range_key)
                                
                    if not is_top_left:
                        continue
                        
                    if is_merged:
                        min_row, min_col, max_row_range, max_col_range = merged_map[(r, c)]
                        lim_max_col = min(max_col_range, max_col)
                        merged_width = sum(col_widths_inches[col_i - 1] for col_i in range(min_col, lim_max_col + 1))
                        word_cell.width = Inches(merged_width)
                    else:
                        word_cell.width = Inches(col_widths_inches[c - 1])
                        
                    val_str = str(cell.value) if cell.value is not None else ""
                    
                    word_cell.text = ""
                    p = word_cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1.15
                    
                    # Alignment mapping
                    if cell.alignment and cell.alignment.horizontal:
                        align = cell.alignment.horizontal
                        if align == "center":
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif align == "right":
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        elif align == "left":
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            
                    run = p.add_run(val_str)
                    run.font.name = 'Arial'
                    
                    if cell.font:
                        if cell.font.name:
                            run.font.name = cell.font.name
                        if cell.font.size:
                            run.font.size = Pt(cell.font.size)
                        else:
                            run.font.size = Pt(10)
                        if cell.font.bold:
                            run.bold = True
                        if cell.font.italic:
                            run.italic = True
                            
                        if cell.font.color and cell.font.color.rgb:
                            f_color = str(cell.font.color.rgb)
                            if len(f_color) == 8:
                                f_color = f_color[2:]
                            if len(f_color) == 6 and re.match(r'^[0-9a-fA-F]+$', f_color):
                                r_c = int(f_color[0:2], 16)
                                g_c = int(f_color[2:4], 16)
                                b_c = int(f_color[4:6], 16)
                                run.font.color.rgb = RGBColor(r_c, g_c, b_c)
                    else:
                        run.font.size = Pt(10)
                        
                    # Shading
                    if has_alert:
                        set_cell_background(word_cell, "FCE4D6")
                    else:
                        bg_color = get_clean_hex_color(cell)
                        if bg_color:
                            set_cell_background(word_cell, bg_color)
                            
            for col_idx, width in enumerate(col_widths_inches):
                table.columns[col_idx].width = Inches(width)
                
        # Chart rendering check
        openpyxl_has_charts = False
        try:
            if hasattr(ws, "_charts") and len(ws._charts) > 0:
                openpyxl_has_charts = True
            elif hasattr(ws, "charts") and len(ws.charts) > 0:
                openpyxl_has_charts = True
        except Exception:
            pass
            
        if openpyxl_has_charts:
            chart_img_stream = None
            try:
                # Load dataframe just for chart parameters
                df_for_chart = pd.read_excel(io.BytesIO(excel_bytes), sheet_name=sheet_name)
                chart_img_stream = generate_matplotlib_chart(df_for_chart, sheet_name)
            except Exception as chart_err:
                print(f"[!] Warning: Matplotlib chart drawing failed: {chart_err}")
                
            if chart_img_stream:
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_p.paragraph_format.space_before = Pt(12)
                img_p.paragraph_format.space_after = Pt(12)
                try:
                    run = img_p.add_run()
                    run.add_picture(chart_img_stream, width=Inches(5.0))
                except Exception as img_err:
                    print(f"[!] Warning: Picture embedding in server engine failed: {img_err}")
            else:
                chart_p = doc.add_paragraph()
                chart_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                chart_p.paragraph_format.space_before = Pt(12)
                chart_p.paragraph_format.space_after = Pt(12)
                
                chart_run = chart_p.add_run(
                    f"[Embedded Workplan Chart: Visual representation of {sheet_name} Data]"
                )
                chart_run.font.name = 'Arial'
                chart_run.font.size = Pt(11)
                chart_run.bold = True
                chart_run.italic = True
                chart_run.font.color.rgb = RGBColor(192, 0, 0)
                
        if index < total_sheets - 1:
            doc.add_page_break()
            
    doc_out = io.BytesIO()
    doc.save(doc_out)
    doc_out.seek(0)
    return doc_out


@app.post("/api/convert-excel-to-word")
async def convert_excel_to_word_endpoint(file: UploadFile = File(...)):
    """
    API endpoint to securely upload an Excel workbook and return a clean, styled,
    cloned Word document, processed 100% in-memory with zero file system writes.
    """
    if not file.filename.endswith('.xlsx'):
        raise HTTPException(status_code=400, detail="Unsupported file format. Please upload a valid Microsoft Excel (.xlsx) file.")
        
    try:
        excel_bytes = await file.read()
        docx_buffer = convert_excel_to_docx_bytes(excel_bytes)
        
        base_name, _ = os.path.splitext(file.filename)
        output_filename = f"{base_name}.docx"
        
        headers = {
            'Content-Disposition': f'attachment; filename="{output_filename}"'
        }
        return StreamingResponse(
            docx_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers
        )
    except ValueError as ve:
        raise HTTPException(status_code=400, detail=str(ve))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Core Processing Pipeline Error: {e}")


# Mount static files for Flutter Web (if build directory exists)
if os.path.exists(os.path.join("build", "web")):
    app.mount("/", StaticFiles(directory="build/web"), name="web")


# =====================================================================
# Main Server Entry Point
# =====================================================================
if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    print(f" [INFO] Starting FastAPI Local Bridge Server on port {port}...")
    uvicorn.run("server:app", host="127.0.0.1", port=port, reload=False)
